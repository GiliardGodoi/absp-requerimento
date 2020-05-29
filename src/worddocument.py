import re
import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.shared import Pt


def salvar_arquivo(worddocument, outputfilename = "RELACAO.docx"):

    if not outputfilename.endswith(".docx"):
        outputfilename += ".docx"

    worddocument.save(outputfilename)

def transform_RelacionarRequerimento(text):
    regex1 = re.compile(r"\s+infra-assinad[aos]{1,2},")
    regex2 = re.compile(r"\s+no uso das atribuições[\s\w]+,")
    regex3 = re.compile(r"\s+requer a Vossa Excelência que seja expedido ofício")
    regex4 = re.compile(r"[\s\n]{2,}")

    text = regex1.sub('', text)
    text = regex2.sub('', text)
    text = regex3.sub('', text)
    text = regex4.sub(' ', text)

    artigo = text[0].lower()

    texto = (', de autoria d'+ artigo + text[1:]).strip()

    if not texto.endswith('.'):
        if texto.endswith(';') or texto.endswith(',') or texto.endswith('-') :
            texto = texto[:len(texto)-1]
        texto += '.'

    return texto


def relacionar_requerimentos(documentos, destino = "."):

    saida = docx.Document()
    saida.core_properties.language = "pt-BR"

    filename = os.path.join(destino, "RELACAO.docx")


    for documento in documentos:
        texto = transform_RelacionarRequerimento(documento["conteudo"])
        nro = documento["nro"]
        tipo = documento["tipo"]
        ano = documento["ano"]

        paragrafo = saida.add_paragraph()
        paragrafo.alignment = ALIGN.JUSTIFY

        run = paragrafo.add_run(f"{tipo} n.º {nro}/{ano}")
        run.bold = True
        run.font.name = "Times New Romam"
        run.font.size = Pt(12)


        run = paragrafo.add_run(texto)
        run.font.name = "Times New Romam"
        run.font.size = Pt(12)


    salvar_arquivo(saida, filename)


def relacionar_requerimentos2(documentos, destino = "."):

    saida = docx.Document()
    saida.core_properties.language = "pt-BR"
    paragrafo = saida.add_paragraph()
    paragrafo.alignment = ALIGN.JUSTIFY

    filename = os.path.join(destino, "RELACAO(2).docx")


    for documento in documentos:
        texto = transform_RelacionarRequerimento(documento["conteudo"])
        nro = documento["nro"]
        tipo = documento["tipo"]
        ano = documento["ano"]

        run = paragrafo.add_run(f" {tipo} n.º {nro}/{ano}")
        run.bold = True
        run.font.name = "Times New Romam"
        run.font.size = Pt(12)


        run = paragrafo.add_run(texto)
        run.font.name = "Times New Romam"
        run.font.size = Pt(12)

        run = paragrafo.add_run(" ")
        run.font.name = "Times New Romam"
        run.font.size = Pt(12)

    salvar_arquivo(saida, filename)